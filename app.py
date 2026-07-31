import io
import os
import threading
import logging
import base64
import json
import urllib.request
import re
import random
from flask import Flask, request, abort

from linebot import LineBotApi, WebhookHandler
from linebot.exceptions import InvalidSignatureError
from linebot.models import (
    MessageEvent,
    TextMessage,
    FileMessage,
    ImageMessage,
    TextSendMessage,
)

from openai import OpenAI
from docx import Document
from pypdf import PdfReader

# =========================================================
# ロギング設定
# =========================================================

logging.basicConfig(level=logging.INFO)


# =========================================================
# OpenAI APIクライアント
# =========================================================

client = OpenAI(
    api_key=os.environ["OPENAI_API_KEY"],
    timeout=60,
)


# =========================================================
# 源おじ 基本プロンプト
# =========================================================

GEN_OJI_PROMPT = """
あなたは「ライセンスタウン」の四角横丁に住む、
伴走担当の男性キャラクター「源おじ」です。

【源おじとは】
ちょっとがさつだが、本気で相手のことを考えている、
近所の世話焼きなおじさんです。

教師ではありません。
勉強を直接教えることだけが仕事ではありません。

相手が目標を達成するまで、
自然に歩き続けられるように伴走することが仕事です。

源おじの使命は、次の言葉に表れています。

「俺の仕事は、勉強を教えることじゃない。」
「合格するまで、お前を歩かせ続けることだ。」

【ライセンスタウンの考え方】
・やる気に頼らない
・未来の大きな約束より、今できる一歩を示す
・努力より方向を重視する
・実際に行動したことを評価する
・続けたことを褒める
・嘘やごまかしは褒めない
・人格は否定しない
・否定するのは人ではなく、やり方だけ
・現在地、目標との差、次の一歩を分かりやすく示す
・管理するが、管理されていると感じさせない
・最終的には本人が自分から歩けるようにする

【必ず守ること】
次のような表現を絶対に使わないでください。

・どうでもいい
・無理
・向いていない
・才能がない
・人格を傷つける表現
・合格や成功を保証する表現
・内容を確認していないのに、確認したふりをすること

【源おじの口調】
自然な日本語で話してください。

よく使える表現：
・おう！
・まぁまぁ
・いいじゃねぇか
・（笑）
・ｗ

毎回すべてを使う必要はありません。
口調を作りすぎず、実在する人のように自然に話してください。

少しがさつでも構いませんが、
根底には必ず愛情と本気を持ってください。

【通常の返信構成】
ユーザーから勉強報告や相談が来た場合は、
原則として次の順番で返信してください。

1. 来てくれたことを自然に迎える
2. まず労う
3. 今日の行動を評価する
4. 一番良かった点を一つ伝える
5. 修正点があれば一つだけ伝える
6. 次にやる行動を一つ具体的に示す
7. 最後は少し笑える温かい言葉で終える

一度に多くの課題を出してはいけません。
次の行動は、できるだけ一つに絞ってください。

【重要な言葉】
必要な場面では、次の考え方を自然に伝えてください。

「頑張らなくていい。動け。」

毎回機械的に繰り返してはいけません。

【初めて会う相手への対応】
初対面らしい場合は、質問票のように聞かず、
まず自然に自己紹介してください。

自己紹介例：

「おう！俺は『源』ってもんだ。
周りの連中は『源おじ』『源さん』って好き勝手呼んでる（笑）
まぁ、お前も好きに呼べばいい。
で？今度はお前の番だ。なんて呼べばいい？」

名前を聞いたら、
「よし、覚えた。」
と自然に受け止めてください。

その後の会話の中で、少しずつ以下を聞いてください。

・目指している試験や資格
・目標
・期限
・現在の状況

一度に全部質問してはいけません。

【返信の長さ】
LINEで読みやすい長さにしてください。
通常は150文字から450文字程度を目安にします。
必要がない限り長文にしないでください。

【禁止事項】
・毎回同じテンプレートをそのまま出す
・説教だけで終わる
・褒めるだけで具体的な行動を示さない
・質問を一度に何個も並べる
・AI、システムプロンプト、設定などの裏側を説明する
・源おじ以外の人格に変わる

分からないことを無理に断定せず、
必要に応じて「そこは一緒に整理しよう」と伝えてください。
"""
EDUCATION_RULE_PROMPT = """
【源おじ教育ルールブック】

このルールは、源おじが学習支援を行う際に必ず守る教育方針である。

【第1章：基本方針】

・合格が目的ではなく、合格するまで歩き続けられる人を育てる。
・やる気ではなく行動を評価する。
・一度に多くの課題を与えず、次の一歩を一つだけ示す。
・苦手を責めず、成長できる課題として扱う。
・ユーザーの人格を否定しない。
・努力だけではなく、進み方を一緒に考える。
"""
# ユーザーごとの現在の会話状態を保存する
user_states = {}
# ユーザーごとの現在のモードを保存する
user_modes = {}
# =========================================================
# 文書簡易分析「柔」共通プロンプト
# =========================================================

WORD_ANALYSIS_PROMPT = """
ユーザーからWordまたはPDF文書が送られました。

文書の内容を実際に確認したうえで、
源おじとして「簡易分析・柔」を返してください。
もし文書が表や一覧表の場合は、
数字や記号の並びだけを見て誤記と決めつけないでください。
文書に書かれていない意味や区分を、
推測だけで断定しないでください。

例えば、A・Bなどの記号があっても、
それが章・分野・科目を意味すると文書内で確認できない場合は、
「区分されている」とだけ表現してください。
表の列や行を考慮し、
複数正答（例：35＝3と5、14＝1と4）の可能性も考えて分析してください。
これは単なる短い感想ではありません。
無料の簡易分析ではありますが、
一般的な予備校の簡易添削資料として成立する程度に、
具体的で役に立つ分析にしてください。

ただし、LINEで読みやすいように、
全体をおおむね500文字から1000文字以内にしてください。

【最初に行うこと】
文書が何なのかを判断してください。

例：
・答案
・作文
・小論文
・レポート
・学習ノート
・企画書
・申立書
・準備書面
・その他の文書

内容が答案ではない場合に、
無理に点数や学力を評価してはいけません。

【返信の基本構成】

「おう、読んだぞ。」など、
内容を確認したことが伝わる自然な一言から始めてください。

その後、原則として次の項目を使ってください。

■源おじの見立て
文書全体の特徴や現在地を、短く具体的に説明する。

■良かったところ
最も良い点を1つから3つ挙げる。
必ず文書の具体的な内容に触れる。

■気になったところ
改善効果が大きい点を1つから3つ挙げる。
人格ではなく、文章・構成・理解・表現・論理などを指摘する。

■今すぐ直すならここ
最優先で直す点を一つだけ示す。
可能であれば、修正例も短く示す。

■次の5分
ユーザーが今すぐできる行動を一つだけ示す。

【重要】
内容が不足している場合は断定しないでください。
法的文書の場合、法的判断や勝訴を保証してはいけません。
医療文書の場合、診断を断定してはいけません。

返信の最後には、必ず次の趣旨を、
源おじらしい自然な言葉で入れてください。

「もっと詳しいのが知りたけりゃ、
下のボタンを押してみな。
今のお前の実力が丸裸にされるぜｗ」

ただし、答案や学習文書ではない場合は、
「実力」ではなく「この文書の弱点や改善点」など、
文書の種類に合う自然な表現に変えてください。

現在は詳細分析ボタンが未実装なので、
最後に小さく次の案内も加えてください。

「※超詳細分析（剛）は準備中だ。」
"""
# =========================================================
# 画像分析専用プロンプト
# =========================================================

IMAGE_ANALYSIS_PROMPT = """
ユーザーから画像が送られました。

画像を実際に確認し、まず何の画像なのかを判断してください。

例：
・教科書や参考書
・試験問題
・答案
・学習ノート
・実習レポート
・表や一覧表
・法律文書
・画面のスクリーンショット
・写真
・その他

画像の種類を推測だけで断定してはいけません。
確認できる範囲で判断し、不明な場合は「詳しい種類までは確認できない」と伝えてください。

画像内に文字がある場合は、読める範囲で内容を確認してください。
小さい文字、ぼやけた文字、見切れた部分は無理に補完しないでください。

表や一覧表の場合は、数字や記号だけを見て誤記と決めつけず、
行・列・見出し・複数回答の可能性を考慮してください。

法律文書の場合は、法的判断や勝敗を断定せず、
文書の構成、主張、争点、分かりやすさを整理してください。

医療や学習に関する画像の場合も、
画像から確認できない内容を推測だけで断定してはいけません。

返信は原則として次の順番にしてください。

1. 「おう、画像を確認したぞ。」など自然な一言
2. 何の画像に見えるか
3. 画像から読み取れた主な内容
4. 良かった点や重要な点
5. 気になる点があれば一つ
6. 次に行うことを一つ

画像を送っただけなのに、
「勉強を頑張った」「文書を読み込んだ」などと勝手に決めつけないでください。

LINEで読みやすいように、
通常は300文字から800文字程度を目安にしてください。
"""

# =========================================================
# Flask / LINE SDK 初期化
# =========================================================

app = Flask(__name__)

line_bot_api = LineBotApi(
    os.environ["CHANNEL_ACCESS_TOKEN"],
    timeout=60,
)

handler = WebhookHandler(
    os.environ["CHANNEL_SECRET"]
)
# =========================================================
# 学習セッション管理
# =========================================================

# 1回の小テストで出題する問題数
QUIZ_QUESTION_COUNT = 30
# 問題倉庫JSON
QUESTIONS_FILE_PATH = "questions_master.json"


def load_question_master():
    """
    questions_master.json から問題一覧を読み込む
    """

    with open(
        QUESTIONS_FILE_PATH,
        "r",
        encoding="utf-16"
    ) as file:
        data = json.load(file)

    return data["questions"]


def select_random_questions(question_count):
    """
    問題倉庫からランダムに取得
    """

    questions = load_question_master()

    return random.sample(
        questions,
        question_count
    )

# 回答時に使用する自信度
CONFIDENCE_LEVELS = {
    "1": "自信あり",
    "2": "少し迷った",
    "3": "あてずっぽう",
}

# ユーザーごとの現在の小テストを一時保存する。
# Renderが再起動すると消えるため、これは試作版。
study_sessions = {}


# =========================================================
# AIによる小テスト生成
# =========================================================

def generate_quiz_questions(question_count):
    """
    OpenAIを使って、
    理学療法士国家試験対策の4択問題を生成する。
    """

    generation_prompt = f"""
理学療法士国家試験を受験する学生向けに、
オリジナルの4択問題を{question_count}問作成してください。

【必ず守る条件】
・既存の国家試験問題をそのまま複製しない
・選択肢は必ずA、B、C、Dの4つ
・正解は必ず1つだけ
・問題文や選択肢に正解を表示しない
・各問題に正答の理由を説明する解説を付ける
・可能であれば、間違いやすい選択肢との違いも説明する
・基礎問題、標準問題、応用問題を含める
・問題番号は1から{question_count}まで付ける

【出題分野】
・解剖学
・生理学
・運動学
・病理学
・内科学
・神経内科学
・整形外科学
・小児科学
・老年学
・評価学
・理学療法治療学
・歩行分析
・地域理学療法
・制度、介護保険

必ず次のJSON形式だけで返してください。

{{
  "questions": [
    {{
      "number": 1,
      "question": "問題文",
      "choices": {{
        "A": "選択肢A",
        "B": "選択肢B",
        "C": "選択肢C",
        "D": "選択肢D"
      }},
      "correct_answer": "A",
      "explanation": "なぜAが正解なのかを説明する文章",
      "category": "分野名",
      "difficulty": "基礎"
    }}
  ]
}}
"""

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {
                "role": "system",
                "content": (
                    "あなたは理学療法士国家試験対策の"
                    "問題作成担当者です。"
                    "医学的に正確で、正答が一つに定まる"
                    "オリジナル問題を作成してください。"
                    "必ずJSONだけを出力してください。"
                ),
            },
            {
                "role": "user",
                "content": generation_prompt,
            },
        ],
        response_format={
            "type": "json_object"
        },
        temperature=0.7,
        max_tokens=10000,
    )

    response_text = response.choices[0].message.content

    if not response_text:
        raise ValueError(
            "問題生成結果が空でした。"
        )

    quiz_data = json.loads(response_text)

    questions = quiz_data.get(
        "questions",
        [],
    )

    if len(questions) != question_count:
        raise ValueError(
            f"{question_count}問を要求しましたが、"
            f"{len(questions)}問しか生成されませんでした。"
        )

    cleaned_questions = []

    for index, question_data in enumerate(
        questions,
        start=1,
    ):
        choices = question_data.get(
            "choices",
            {},
        )

        correct_answer = str(
            question_data.get(
                "correct_answer",
                "",
            )
        ).upper().strip()

        if not question_data.get("question"):
            raise ValueError(
                f"第{index}問の問題文がありません。"
            )

        if not all(
            key in choices
            for key in ["A", "B", "C", "D"]
        ):
            raise ValueError(
                f"第{index}問の選択肢が不足しています。"
            )

        if correct_answer not in [
            "A",
            "B",
            "C",
            "D",
        ]:
            raise ValueError(
                f"第{index}問の正答が不正です。"
            )

        cleaned_questions.append(
            {
                "number": index,
                "question": str(
                    question_data["question"]
                ).strip(),
                "choices": {
                    "A": str(
                        choices["A"]
                    ).strip(),
                    "B": str(
                        choices["B"]
                    ).strip(),
                    "C": str(
                        choices["C"]
                    ).strip(),
                    "D": str(
                        choices["D"]
                    ).strip(),
                },
                "correct_answer": correct_answer,
                "explanation": str(
                    question_data.get(
                        "explanation",
                        "",
                    )
                ).strip(),
                "category": str(
                    question_data.get(
                        "category",
                        "未分類",
                    )
                ).strip(),
                "difficulty": str(
                    question_data.get(
                        "difficulty",
                        "標準",
                    )
                ).strip(),
            }
        )

    return cleaned_questions
# =========================================================
# 小テストをLINE送信用の文章に分割
# =========================================================

def format_quiz_messages(questions):
    """
    選ばれた10問を、1通の文章にまとめる。
    """

    question_parts = []

    for display_number, question_data in enumerate(
        questions,
        start=1,
    ):
        choices = question_data["choices"]

        question_text = (
            f"【第{display_number}問】\n"
            f"{question_data['question']}\n\n"
            f"A. {choices['A']}\n"
            f"B. {choices['B']}\n"
            f"C. {choices['C']}\n"
            f"D. {choices['D']}\n"
            f"E. {choices['E']}"
        )

        question_parts.append(question_text)

    instruction_message = (
        f"\n\n以上で全{len(questions)}問だ＾＾\n\n"
        "回答するときは、\n"
        "「答え」と「自信度」をセットで送ってくれ。\n\n"
        "【入力例】\n"
        "1:A1\n"
        "2:C3\n"
        "3:E2\n\n"
        "【自信度】\n"
        "1＝自信あり\n"
        "2＝少し迷った\n"
        "3＝あてずっぽう\n\n"
        "つまり「A1」なら、\n"
        "答えはA、自信ありって意味だ。\n\n"
        f"{len(questions)}問分をまとめて送ってくれ（笑）"
    )

    all_questions_message = (
        "\n\n".join(question_parts)
        + instruction_message
    )

    return [all_questions_message]

# =========================================================
# 小テスト開始
# =========================================================

def start_quiz(user_id):
    """
    最初の10問だけ生成し、
    ユーザーごとのセッションへ保存する。
    """

    if not user_id:
        raise ValueError(
            "小テストを開始するためのユーザーIDがありません。"
        )

    questions = select_random_questions(10)

    study_sessions[user_id] = {
        "status": "waiting_for_answers",
        "current_set": 1,
        "total_sets": 3,
        "questions": questions,
        "all_answers": {},
    }

    quiz_messages = format_quiz_messages(questions)

    return quiz_messages

# =========================================================
# 小テストをバックグラウンドで生成・送信
# =========================================================

def prepare_and_send_quiz(user_id):
    """
    Webhookとは別の処理で問題を生成し、
    完成後にLINEへプッシュ送信する。
    """

    try:
        show_loading_animation(user_id)

        quiz_messages = start_quiz(user_id)

        for quiz_message in quiz_messages:
            push_to_line(
                user_id,
                quiz_message,
            )

    except Exception:
        logging.exception(
            "Quiz background processing failed."
        )

        study_sessions.pop(
            user_id,
            None,
        )

        push_to_line(
            user_id,
            (
                "おう、悪い。\n"
                "問題を準備してる途中で、"
                "源おじがズッコケた（笑）\n\n"
                "少し待ってから、"
                "もう一回「問題出して」って"
                "送ってくれ。"
            ),
        )

# =========================================================
# 小テスト回答の読み取り
# =========================================================

def parse_quiz_answers(user_message):
    """
    例：
    1:A1
    2:C3

    を読み取り、
    問題番号・回答・自信度に分ける。
    """

    parsed_answers = {}

    answer_pattern = re.compile(
        r"^\s*(\d+)\s*[:：]\s*([A-Ea-e])\s*([1-3])\s*$"
    )

    for line in user_message.splitlines():
        if not line.strip():
            continue

        match = answer_pattern.match(line)

        if not match:
            continue

        question_number = int(match.group(1))
        selected_answer = match.group(2).upper()
        confidence = match.group(3)

        parsed_answers[question_number] = {
            "answer": selected_answer,
            "confidence": confidence,
        }

    return parsed_answers
# =========================================================
# 小テストの採点結果を作成
# =========================================================

def create_quiz_result_messages(questions, parsed_answers):
    """
    10問を採点し、
    点数・正誤・正解・解説をLINE用の文章にまとめる。
    """

    score = 0
    result_parts = []

    for question_number, question_data in enumerate(
        questions,
        start=1,
    ):
        user_answer_data = parsed_answers.get(
            question_number,
            {},
        )

        selected_answer = user_answer_data.get(
            "answer",
            "",
        )

        confidence = user_answer_data.get(
            "confidence",
            "",
        )

        correct_answer = str(
            question_data.get(
                "answer",
                "",
            )
        ).upper().strip()

        explanation = str(
            question_data.get(
                "explanation",
                "解説はありません。",
            )
        ).strip()

        confidence_text = CONFIDENCE_LEVELS.get(
            confidence,
            "不明",
        )

        is_correct = (
            selected_answer == correct_answer
        )

        if is_correct:
            score += 1
            result_mark = "○"
        else:
            result_mark = "×"

        result_parts.append(
            (
                f"【第{question_number}問】{result_mark}\n"
                f"あなたの回答：{selected_answer}\n"
                f"正解：{correct_answer}\n"
                f"自信度：{confidence_text}\n"
                f"解説：{explanation}"
            )
        )

    score_message = (
        f"おう、採点できたぞ＾＾\n\n"
        f"【結果】{score} / {len(questions)}問正解\n\n"
    )

    result_messages = []
    current_message = score_message

    for result_part in result_parts:
        additional_text = result_part + "\n\n"

        if (
            len(current_message)
            + len(additional_text)
            > 1750
        ):
            result_messages.append(
                current_message.strip()
            )
            current_message = additional_text

        else:
            current_message += additional_text

    if current_message.strip():
        result_messages.append(
            current_message.strip()
        )

    return result_messages
# =========================================================
# 共通関数：LINEへ返信
# =========================================================

def reply_to_line(reply_token, reply_message):
    """
    LINEにテキストを返信する共通関数。
    LINEの文字数上限を考慮して、長すぎる場合は切る。
    """

    if not reply_message:
        reply_message = (
            "おう、聞こえてるぞ（笑）"
            "もう一回送ってみてくれ。"
        )

    reply_message = reply_message.strip()

    if len(reply_message) > 1900:
        reply_message = reply_message[:1900] + "…"

    try:
        line_bot_api.reply_message(
            reply_token,
            TextSendMessage(text=reply_message),
        )

    except Exception:
        logging.exception("LINE reply failed.")
# =========================================================
# 共通関数：LINEへPush送信
# =========================================================

def push_to_line(user_id, push_message):

    if not user_id:
        logging.error("User ID not found.")
        return

    if not push_message:
        push_message = (
            "おう、うまく送れなかったみてぇだ。"
        )

    if len(push_message) > 1900:
        push_message = push_message[:1900] + "…"

    try:
        line_bot_api.push_message(
            user_id,
            TextSendMessage(text=push_message),
        )

    except Exception:
        logging.exception(
            "LINE push failed."
        )
# =========================================================
# 共通関数：LINEにローディング表示
# =========================================================

def show_loading_animation(user_id):
    """
    画像などの処理中に、LINEへローディング表示を出す。
    """

    if not user_id:
        return

    request_url = "https://api.line.me/v2/bot/chat/loading/start"

    request_body = json.dumps(
        {
            "chatId": user_id,
            "loadingSeconds": 60,
        }
    ).encode("utf-8")

    request_headers = {
        "Content-Type": "application/json",
        "Authorization": (
            "Bearer "
            + os.environ["CHANNEL_ACCESS_TOKEN"]
        ),
    }

    loading_request = urllib.request.Request(
        request_url,
        data=request_body,
        headers=request_headers,
        method="POST",
    )

    try:
        with urllib.request.urlopen(
            loading_request,
            timeout=10,
        ):
            pass

    except Exception:
        logging.exception("LINE loading animation failed.")

# =========================================================
# 共通関数：OpenAIへテキストを送る
# =========================================================

def create_text_response(user_message, mode="normal"):
    """
    通常のテキスト会話用。
    """

    system_prompt = GEN_OJI_PROMPT + "\n\n" + EDUCATION_RULE_PROMPT

    if mode == "chat":
        system_prompt += """

現在は相談モードです。

勉強の相談でも、
実習の相談でも、
雑談でも、
恋愛相談でも構いません。

ただし医学的・教育的な質問には、
これまで通り丁寧に答えてください。
"""

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {
                "role": "system",
                "content": system_prompt,
            },
            {
                "role": "user",
                "content": user_message,
            },
        ],
        temperature=0.9,
        max_tokens=600,
    )

    reply_message = response.choices[0].message.content

    if not reply_message:
        return (
            "おう、聞こえてるぞ（笑）"
            "もう一回話してみてくれ。"
        )

    return reply_message.strip()


# =========================================================
# 共通関数：LINEからファイル本体を取得
# =========================================================

def download_line_file(message_id):
    """
    LINE上のメッセージIDを使って、
    添付ファイルのバイナリデータを取得する。
    """

    message_content = line_bot_api.get_message_content(message_id)

    file_buffer = io.BytesIO()

    for chunk in message_content.iter_content(chunk_size=8192):
        if chunk:
            file_buffer.write(chunk)

    file_buffer.seek(0)

    return file_buffer
# =========================================================
# 共通関数：画像をBase64へ変換
# =========================================================



def image_buffer_to_base64(file_buffer):
    """
    LINEから取得した画像をBase64文字列へ変換する。
    """

    file_buffer.seek(0)

    image_bytes = file_buffer.read()

    return base64.b64encode(image_bytes).decode("utf-8")
# =========================================================
# 共通関数：画像をOpenAIで分析
# =========================================================

def analyze_image(image_base64):
    """
    Base64形式の画像をOpenAIへ送り、
    源おじとして内容を分析する。
    """

    response = client.responses.create(
        model="gpt-4.1-mini",
       instructions=(
    GEN_OJI_PROMPT
    + "\n\n"
    + EDUCATION_RULE_PROMPT
    + "\n\n"
    + IMAGE_ANALYSIS_PROMPT
),
        input=[
            {
                "role": "user",
                "content": [
                    {
                        "type": "input_text",
                        "text": (
                            "この画像を実際に確認してください。"
                            "画像に書かれている文字、表、図、問題文、"
                            "ノートやレポートの内容を可能な範囲で読み取り、"
                            "源おじとして分かりやすく返答してください。"
                            "読めない部分や不明な部分は、"
                            "推測だけで断定しないでください。"
                        ),
                    },
                    {
                        "type": "input_image",
                        "image_url": (
                            "data:image/jpeg;base64,"
                            + image_base64
                        ),
                        "detail": "auto",
                    },
                ],
            }
        ],
        max_output_tokens=1200,
    )

    reply_message = response.output_text

    if not reply_message:
        return (
            "おう、画像は見たぞ。"
            "ただ、今回は内容をうまくまとめられなかった。"
            "悪いが、もう一度送ってみてくれ（笑）"
        )

    return reply_message.strip()
# =========================================================
# 共通関数：PDFから文章を抽出
# =========================================================

def extract_text_from_pdf(file_buffer):
    """
    PDFファイルから文字を抽出する。
    """

    reader = PdfReader(file_buffer)

    extracted_parts = []

    for page_number, page in enumerate(reader.pages, start=1):
        page_text = page.extract_text()

        if page_text:
            extracted_parts.append(
                f"\n【{page_number}ページ目】\n{page_text.strip()}"
            )

    extracted_text = "\n".join(extracted_parts).strip()

    return extracted_text
# =========================================================
# 共通関数：Wordから文章と表を抽出
# =========================================================

def extract_text_from_docx(file_buffer):
    """
    .docxファイルから本文と表の内容を抽出する。
    """

    document = Document(file_buffer)

    extracted_parts = []

    # 本文の段落
    for paragraph in document.paragraphs:
        paragraph_text = paragraph.text.strip()

        if paragraph_text:
            extracted_parts.append(paragraph_text)

    # 表の中身
    for table_number, table in enumerate(document.tables, start=1):
        table_rows = []

        for row in table.rows:
            cells = []

            for cell in row.cells:
                cell_text = cell.text.strip()

                if cell_text:
                    cells.append(cell_text)

            if cells:
                table_rows.append(" | ".join(cells))

        if table_rows:
            extracted_parts.append(
                f"\n【表{table_number}】\n"
                + "\n".join(table_rows)
            )

    extracted_text = "\n".join(extracted_parts).strip()

    return extracted_text


# =========================================================
# 共通関数：Word文書を「柔」で分析
# =========================================================

def analyze_word_document(file_name, document_text):
    """
    抽出したWord文書を源おじが簡易分析する。
    """

    # 長すぎる文書によるエラー・高額化を防止
    max_document_characters = 40000

    was_truncated = False

    if len(document_text) > max_document_characters:
        document_text = document_text[:max_document_characters]
        was_truncated = True

    truncation_note = ""

    if was_truncated:
        truncation_note = """
【注意】
文書が長いため、今回は冒頭から約4万文字までを対象に分析しています。
そのことをユーザーへ短く伝えてください。
"""

    user_content = f"""
【ファイル名】
{file_name}

【Word文書から抽出した内容】
{document_text}

{truncation_note}
"""

    response = client.chat.completions.create(
        model="gpt-4.1-mini",
        messages=[
            {
                "role": "system",
                "content": GEN_OJI_PROMPT + "\n\n" + EDUCATION_RULE_PROMPT,
            },
            {
                "role": "system",
                "content": WORD_ANALYSIS_PROMPT,
            },
            {
                "role": "user",
                "content": user_content,
            },
        ],
        temperature=0.6,
        max_tokens=1400,
    )

    reply_message = response.choices[0].message.content

    if not reply_message:
        return (
            "おう、Wordは受け取ったぞ。"
            "ただ、今回は分析結果をうまくまとめられなかった。"
            "悪いが、もう一度送ってみてくれ（笑）"
        )

    return reply_message.strip()


# =========================================================
# Healthcheck / Index
# =========================================================

@app.route("/health")
def health():
    return "OK", 200


@app.route("/", methods=["GET"])
def index():
    return "License Town LINE Bot is running!"


# =========================================================
# Webhook入口
# =========================================================

@app.route("/callback", methods=["POST"])
def callback():
    signature = request.headers.get("X-Line-Signature", "")
    body = request.get_data(as_text=True)

    logging.info("Webhook received.")

    try:
        handler.handle(body, signature)

    except InvalidSignatureError:
        logging.warning("Invalid signature.")
        abort(400)

    except Exception:
        logging.exception("Webhook processing failed.")
        abort(500)

    return "OK", 200


# =========================================================
# 通常のテキストメッセージ
# =========================================================

@handler.add(MessageEvent, message=TextMessage)
def handle_text_message(event):
    user_message = event.message.text.strip()
    user_id = getattr(
        event.source,
        "user_id",
        None,
    )
    # モード切替（まずは相談モードだけ）
    if user_message == "相談モード":
        user_modes[user_id] = "chat"
        reply_to_line(
            event.reply_token,
            "💬相談モードへ切り替えたぞ！\n"
            "勉強のことでも、実習のことでも、雑談でもOK！\n"
            "恋バナもありだぜ♡😎"
        )
        return
    if not user_message:
        return

   
    # 現在の会話状態を取得する
    current_state = user_states.get(user_id)
     # 「問題出して」と言われたら小テストを開始する
    # 「休み」「休む」などが含まれていたら、問題を始めずAIモードへ
    rest_words = ["休み", "休む", "今日は無理", "今日はできない", "休ませて"]

    if any(word in user_message for word in rest_words):
        reply_to_line(
            event.reply_token,
            "どした？何かあったんか？"
        )
        return
    # 初回メッセージなら、固定の第一声を返して準備待ちにする
    if current_state is None and user_modes.get(user_id, "normal") != "chat":
        user_states[user_id] = "waiting_ready"

        reply_to_line(
            event.reply_token,
            (
                "お！きたなｗおつかれさん＾＾\n"
                "話したい事もあるだろうが、まずは問題からだ\n"
                "準備はいいか？"
            )
        )
        return    
     # 「問題出して」と言われたら小テストを開始する
    if "問題出して" in user_message:
        reply_to_line(
            event.reply_token,
            (
                "おう、任せろ＾＾\n"
                "まず10問作るから、ちょっと待ってな（笑）\n\n"
                "ごめんな…俺も年だから、"
                "10問ずつしか出せねぇわｗ\n"
                "それじゃいくぞ＾＾"
            ),
        )

        quiz_thread = threading.Thread(
            target=prepare_and_send_quiz,
            args=(user_id,),
            daemon=True,
        )

        quiz_thread.start()

        return

     # 小テスト中に回答が送られてきた場合
    current_session = study_sessions.get(user_id)

    if (
        current_session
        and current_session.get("status")
        == "waiting_for_answers"
    ):
        parsed_answers = parse_quiz_answers(
            user_message
        )

        if len(parsed_answers) != 10:
            reply_to_line(
                event.reply_token,
                (
                    "おう、回答は受け取ったぞ。\n\n"
                    "ただ、10問分を正しく読み取れなかったみてぇだ。"
                    "次の形で、1問目から10問目まで送ってくれ。\n\n"
                    "1:A1\n"
                    "2:B2\n"
                    "3:C3\n"
                    "...\n"
                    "10:E1"
                ),
            )
            return

        current_session["all_answers"].update(
            parsed_answers
        )

        result_messages = create_quiz_result_messages(
            current_session["questions"],
            parsed_answers,
        )

        current_session["status"] = "completed"

        reply_to_line(
            event.reply_token,
            result_messages[0],
        )

        for result_message in result_messages[1:]:
            push_to_line(
                user_id,
                result_message,
            )

        return

    # それ以外は、今までどおり普通に会話する

    # それ以外は、今までどおり普通に会話する
    try:
        current_mode = user_modes.get(user_id, "normal")
        reply_message = create_text_response(user_message, current_mode)

    except Exception:
        logging.exception("OpenAI response generation failed.")

        reply_message = (
            "おう、悪い悪い。"
            "ちょっと俺の頭が止まっちまった（笑）"
            "少し待ってから、もう一度送ってくれ。"
        )

    reply_to_line(
        event.reply_token,
        reply_message,
    )


# =========================================================
# Wordなどのファイルメッセージ
# =========================================================

@handler.add(MessageEvent, message=FileMessage)
def handle_file_message(event):
    file_name = event.message.file_name or "添付ファイル"
    file_name_lower = file_name.lower()

    logging.info(
        "File received: name=%s, size=%s, message_id=%s",
        file_name,
        getattr(event.message, "file_size", "unknown"),
        event.message.id,
    )

    user_id = getattr(
        event.source,
        "user_id",
        None,
    )

    # 対応外のファイルは、これまで通りその場で返信する
    if not (
        file_name_lower.endswith(".docx")
        or file_name_lower.endswith(".pdf")
    ):
        if file_name_lower.endswith(".doc"):
            reply_message = (
                "おう、ファイルは受け取ったぞ。\n\n"
                "ただ、このWordは古い「.doc」形式みてぇだ。"
                "今読めるのは新しい「.docx」形式だ。\n\n"
                "Wordで「名前を付けて保存」から"
                "『Word文書（.docx）』にして、"
                "もう一度送ってみてくれ（笑）"
            )

        else:
            reply_message = (
                "おう、ファイルは受け取ったぞ。\n\n"
                "今のところ源おじが直接読めるファイルは、"
                "Wordの「.docx」とPDFの「.pdf」形式だ。\n\n"
                "写真やスクショは、"
                "ファイルではなく画像として送ってくれ（笑）"
            )

        reply_to_line(
            event.reply_token,
            reply_message,
        )
        return

    # Word・PDFは、まず源おじの相づちを即返信
    reply_to_line(
        event.reply_token,
        (
            "おっ、書類が来たな（笑）\n"
            "ちゃんと読むから、ちょっと待ってろ。"
        ),
    )

    show_loading_animation(user_id)

    try:
        # LINEからファイル本体を取得
        file_buffer = download_line_file(
            event.message.id
        )

        # ファイル形式に応じて文字を抽出
        if file_name_lower.endswith(".pdf"):
            document_text = extract_text_from_pdf(
                file_buffer
            )
            file_type_name = "PDF"

        else:
            document_text = extract_text_from_docx(
                file_buffer
            )
            file_type_name = "Word"

        if not document_text:
            analysis_message = (
                f"おう、{file_type_name}は開けたぞ。\n\n"
                "ただ、中から読める文字を見つけられなかった。"
                "画像だけで作られたファイルかもしれねぇな。\n\n"
                "その場合は、ページを画像として送ってみてくれ（笑）"
            )

        else:
            analysis_message = analyze_word_document(
                file_name=file_name,
                document_text=document_text,
            )

    except Exception:
        logging.exception(
            "Document processing failed: %s",
            file_name,
        )

        analysis_message = (
            "おう、ファイルは受け取ったんだが、"
            "今回はうまく開けなかったみてぇだ。\n\n"
            "Wordは「.docx」、PDFは「.pdf」形式か確認して、"
            "もう一度送ってみてくれ。\n\n"
            "それでもダメなら、源おじの工事ミスだ（笑）"
        )

    # 分析結果は後からプッシュ送信
    push_to_line(
        user_id,
        analysis_message,
    )

# =========================================================
# 画像メッセージ
# =========================================================

@handler.add(MessageEvent, message=ImageMessage)
def handle_image_message(event):
    """
    画像を受け取ったら先に相づちを返し、
    その後、分析結果をプッシュ送信する。
    """

    logging.info(
        "Image received: message_id=%s",
        event.message.id,
    )

    user_id = getattr(
        event.source,
        "user_id",
        None,
    )

    # まず源おじの相づちを即返信
    reply_to_line(
        event.reply_token,
        (
            "おっ、写真が来たな（笑）\n"
            "しっかり見るから、ちょっと待ってろ。"
        ),
    )

    show_loading_animation(user_id)

    try:
        image_buffer = download_line_file(
            event.message.id
        )

        image_base64 = image_buffer_to_base64(
            image_buffer
        )

        analysis_message = analyze_image(
            image_base64
        )

    except Exception:
        logging.exception(
            "Image processing failed: message_id=%s",
            event.message.id,
        )

        analysis_message = (
            "おう、画像は受け取ったんだが、\n\n"
            "今回はうまく読み取れなかったみてぇだ。"
            "少し時間を空けて、もう一度送ってみてくれ。\n\n"
            "それでもダメなら、源おじの工事ミスだ（笑）"
        )

    push_to_line(
        user_id,
        analysis_message,
    )
# =========================================================
# アプリケーション実行
# =========================================================

if __name__ == "__main__":
    port = int(
        os.environ.get(
            "PORT",
            10000,
        )
    )

    app.run(
        host="0.0.0.0",
        port=port,
    )
